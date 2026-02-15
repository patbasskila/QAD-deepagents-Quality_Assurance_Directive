import json
import logging
import os
import re
import uuid
from dataclasses import asdict
from typing import List, Tuple, Dict, Any


from app.services.types import DocumentBlock
from app.utils.config import get_settings
from app.utils.files import ensure_dir, write_text

log = logging.getLogger("app.ingest")


def _new_id() -> str:
    return str(uuid.uuid4())


def _normalize_ws(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _basic_ocr_quality(text: str) -> Dict[str, Any]:
    """
    Lightweight, deterministic OCR quality heuristics.
    Returns metrics + a suggested needs_review boolean.
    """
    t = text or ""
    n_chars = len(t)
    n_alpha = sum(ch.isalpha() for ch in t)
    n_digit = sum(ch.isdigit() for ch in t)
    n_space = sum(ch.isspace() for ch in t)
    n_printable = sum(ch.isprintable() for ch in t)

    # "weird" = non-ascii printable characters (often OCR artifacts)
    n_weird = sum((ord(ch) > 127) and ch.isprintable() for ch in t)

    # crude word count
    words = re.findall(r"[A-Za-z0-9]+", t)
    n_words = len(words)

    alpha_ratio = (n_alpha / n_chars) if n_chars else 0.0
    printable_ratio = (n_printable / n_chars) if n_chars else 0.0
    weird_ratio = (n_weird / n_chars) if n_chars else 0.0
    avg_word_len = (sum(len(w) for w in words) / n_words) if n_words else 0.0

    # Heuristic thresholds (tune later)
    # - Very low alpha ratio often means garbage OCR or mostly line noise
    # - High weird ratio can indicate encoding / OCR corruption
    # - Extremely low word count suggests OCR didn't extract meaningful text
    needs_review = (
        n_words < 80
        or alpha_ratio < 0.55
        or printable_ratio < 0.95
        or weird_ratio > 0.02
        or avg_word_len < 3.2
    )

    # Simple quality label
    if n_words >= 200 and alpha_ratio >= 0.65 and weird_ratio <= 0.01:
        quality = "good"
    elif n_words >= 120 and alpha_ratio >= 0.58 and weird_ratio <= 0.02:
        quality = "ok"
    else:
        quality = "poor"

    return {
        "n_chars": n_chars,
        "n_words": n_words,
        "alpha_ratio": round(alpha_ratio, 4),
        "printable_ratio": round(printable_ratio, 4),
        "weird_ratio": round(weird_ratio, 4),
        "avg_word_len": round(avg_word_len, 2),
        "quality": quality,
        "needs_review": needs_review,
    }


def _aggregate_ocr_quality(blocks: List["DocumentBlock"]) -> Dict[str, Any]:
    """
    Aggregate OCR quality across OCR blocks.
    """
    ocr_texts = [b.text for b in blocks if (b.block_kind == "ocr_page" and b.text)]
    combined = "\n".join(ocr_texts).strip()

    if not combined:
        return {
            "has_ocr_text": False,
            "quality": "none",
            "needs_review": True,
            "metrics": {"n_chars": 0, "n_words": 0, "alpha_ratio": 0, "weird_ratio": 0, "avg_word_len": 0},
        }

    metrics = _basic_ocr_quality(combined)
    return {
        "has_ocr_text": True,
        "quality": metrics["quality"],
        "needs_review": metrics["needs_review"],
        "metrics": metrics,
    }


def _flatten_blocks(blocks: List[DocumentBlock]) -> str:
    """
    Produce a single text file used for later chunking.
    Keep page/table hints in-line lightly.
    """
    parts: List[str] = []
    for b in blocks:
        header = []
        if b.source_type == "pdf" and b.page_num is not None:
            header.append(f"[PDF page {b.page_num}]")
        if b.block_kind:
            header.append(f"[{b.block_kind}]")
        prefix = " ".join(header)
        if prefix:
            parts.append(prefix)
        parts.append(b.text)
        parts.append("")  # blank line between blocks
    return _normalize_ws("\n".join(parts))


def _write_json(path: str, payload: Any) -> None:
    ensure_dir(os.path.dirname(path))
    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2, ensure_ascii=False)


def ingest_contract_to_blocks(input_path: str, original_name: str) -> Tuple[List[DocumentBlock], Dict[str, Any]]:
    """
    Main entry point for Section 2.
    Returns:
      - blocks: normalized DocumentBlocks
      - report: dict with stats (pages, used_ocr, etc.)
    """
    ext = os.path.splitext(original_name)[1].lower()
    if ext == ".pdf":
        return _ingest_pdf(input_path, original_name)
    if ext == ".docx":
        return _ingest_docx(input_path, original_name)
    raise ValueError("Unsupported file type. Only .pdf and .docx are supported.")


def _ingest_pdf(pdf_path: str, source_name: str) -> Tuple[List[DocumentBlock], Dict[str, Any]]:
    settings = get_settings()
    blocks: List[DocumentBlock] = []
    used_ocr = False

    # --- Primary: extract embedded text via pypdf ---
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as e:
        raise RuntimeError(f"Missing dependency pypdf: {e}")

    reader = PdfReader(pdf_path)
    num_pages = len(reader.pages)

    extracted_chars = 0
    per_page_chars: List[int] = []

    for i, page in enumerate(reader.pages, start=1):
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        txt = _normalize_ws(txt)
        per_page_chars.append(len(txt))
        extracted_chars += len(txt)

        # Create a page block even if empty; OCR may fill later
        if txt:
            blocks.append(
                DocumentBlock(
                    id=_new_id(),
                    text=txt,
                    source_type="pdf",
                    source_name=source_name,
                    page_num=i,
                    block_kind="pdf_text",
                    metadata={"page": i},
                )
            )

    # Heuristic: if most pages have little/no text, OCR fallback
    ocr_enabled = settings.ocr_enabled
    low_text_pages = sum(1 for c in per_page_chars if c < 40)
    should_ocr = ocr_enabled and (extracted_chars < 300 or low_text_pages >= max(1, int(0.6 * num_pages)))

    ocr_pages_processed = 0

    if should_ocr:
        used_ocr = True
        ocr_blocks = _ocr_pdf_pages(pdf_path, source_name)
        blocks.extend(ocr_blocks)
        ocr_pages_processed = len({b.page_num for b in ocr_blocks if b.page_num is not None})

    # Ensure stable ordering: by page_num then block_kind
    blocks.sort(key=lambda b: (b.page_num or 10**9, b.block_kind or ""))

    ocr_quality = _aggregate_ocr_quality(blocks)

    report = {
        "source_type": "pdf",
        "source_name": source_name,
        "num_pages": num_pages,
        "embedded_text_chars": extracted_chars,
        "low_text_pages": low_text_pages,
        "ocr_enabled": ocr_enabled,
        "used_ocr": used_ocr,
        "ocr_pages_processed": ocr_pages_processed,
        "blocks_count": len(blocks),
        "ocr_quality": ocr_quality,
    }
    return blocks, report


def _ocr_pdf_pages(pdf_path: str, source_name: str) -> List[DocumentBlock]:
    """
    OCR fallback for scanned PDFs:
    - convert each page to image (pdf2image)
    - run pytesseract on the image
    """
    settings = get_settings()

    try:
        import pytesseract  # type: ignore
    except Exception as e:
        raise RuntimeError(f"OCR enabled but pytesseract is missing: {e}")

    # Allow explicit tesseract path via env
    if settings.tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

    images = None

    # Strategy A (preferred if available): PyMuPDF render (pip-only; no Poppler)
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(pdf_path)
        dpi = 350
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        rendered = []
        for page in doc:
            pix = page.get_pixmap(matrix=mat, alpha=False)
            # Convert pixmap -> PIL Image
            from PIL import Image

            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            rendered.append(img)
        images = rendered
    except Exception:
        images = None

    # Strategy B: pdf2image (requires Poppler). Keep as fallback for environments that have it.
    if images is None:
        try:
            from pdf2image import convert_from_path  # type: ignore

            images = convert_from_path(pdf_path, dpi=350)
        except Exception as e:
            raise RuntimeError(
                "Failed converting PDF pages to images for OCR. "
                "Install PyMuPDF (recommended) or pdf2image + Poppler.\n"
                f"Error: {e}"
            )

    blocks: List[DocumentBlock] = []

    from PIL import ImageOps

    tess_config = "--oem 3 --psm 6"

    for idx, img in enumerate(images, start=1):
        try:
            gray = ImageOps.grayscale(img)
            gray = ImageOps.autocontrast(gray)
            bw = gray.point(lambda x: 0 if x < 160 else 255, mode="1")

            text = pytesseract.image_to_string(bw, lang="eng", config=tess_config)
        except Exception as e:
            log.warning("OCR failed on page %s: %s", idx, e)
            text = ""

        text = _normalize_ws(text)
        if not text:
            continue

        blocks.append(
            DocumentBlock(
                id=_new_id(),
                text=text,
                source_type="pdf",
                source_name=source_name,
                page_num=idx,
                block_kind="ocr_page",
                metadata={"page": idx, "dpi": 350},
            )
        )

    return blocks


def _ingest_docx(docx_path: str, source_name: str) -> Tuple[List[DocumentBlock], Dict[str, Any]]:
    try:
        import docx  # python-docx
    except Exception as e:
        raise RuntimeError(f"Missing dependency python-docx: {e}")

    d = docx.Document(docx_path)
    blocks: List[DocumentBlock] = []

    para_count = 0
    table_cell_count = 0
    chars = 0

    # Paragraphs
    for p in d.paragraphs:
        txt = _normalize_ws(p.text or "")
        if not txt:
            continue
        para_count += 1
        chars += len(txt)
        blocks.append(
            DocumentBlock(
                id=_new_id(),
                text=txt,
                source_type="docx",
                source_name=source_name,
                page_num=None,
                block_kind="paragraph",
                metadata={},
            )
        )

    # Tables (cells)
    for t_idx, table in enumerate(d.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_txt = _normalize_ws(cell.text or "")
                if not cell_txt:
                    continue
                table_cell_count += 1
                chars += len(cell_txt)
                blocks.append(
                    DocumentBlock(
                        id=_new_id(),
                        text=cell_txt,
                        source_type="docx",
                        source_name=source_name,
                        page_num=None,
                        block_kind="table_cell",
                        metadata={"table": t_idx, "row": r_idx, "col": c_idx},
                    )
                )

    report = {
        "source_type": "docx",
        "source_name": source_name,
        "paragraph_blocks": para_count,
        "table_cell_blocks": table_cell_count,
        "extracted_chars": chars,
        "blocks_count": len(blocks),
    }
    return blocks, report


def persist_ingest_artifacts(work_dir: str, blocks: List[DocumentBlock], report: Dict[str, Any]) -> Dict[str, str]:
    """
    Writes ingestion outputs into <work_dir>/artifacts and returns artifact paths.
    """
    artifacts_dir = os.path.join(work_dir, "artifacts")
    ensure_dir(artifacts_dir)

    blocks_path = os.path.join(artifacts_dir, "document_blocks.json")
    text_path = os.path.join(artifacts_dir, "contract_text.txt")
    report_path = os.path.join(artifacts_dir, "ingest_report.json")

    # JSON: blocks
    _write_json(blocks_path, [asdict(b) for b in blocks])

    # Text: flattened blocks
    flat = _flatten_blocks(blocks)
    write_text(text_path, flat)

    # JSON: report
    _write_json(report_path, report)

    return {
        "document_blocks": blocks_path,
        "contract_text": text_path,
        "ingest_report": report_path,
    }
